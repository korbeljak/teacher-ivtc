from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Tuple
import random

from docx import Document
import openpyxl as op


class ExcelToTest:

    DB_SHEET_NAME: str = "DB"
    CFG_SHEET_NAME: str = "Configuration"

    ANIMALS = ["unicorns", "dragons", "ponies",
               "griffins", "hippos", "otters", "bears"]

    LETTERS = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

    db_file: Path
    variant_type: str
    variant_cnt: int
    db_val_row_cnt: int
    output_dir: Path
    headings = Tuple[str]
    headings_cnt: int

    def __init__(self,
                 db_file: Path,
                 variant_type: str,
                 variant_cnt: int,
                 db_val_row_cnt: int,
                 output_dir: Path):

        self.time_of_gen = datetime.now()

        self.db_file = db_file
        assert db_file.is_file()

        self.variant_type = variant_type
        assert variant_type in ["Animals", "Numbers", "Letters"]

        self.variant_cnt = variant_cnt
        if variant_type == "Animals":
            assert variant_cnt <= len(self.ANIMALS)

        if variant_type == "Letters":
            assert variant_cnt <= len(self.LETTERS)

        self.db_val_row_cnt = db_val_row_cnt

        self.debug = False

        self.output_dir = output_dir
        assert output_dir.is_dir()

    def enable_debugging(self):
        self.debug = True

    def _l(self, msg):
        if self.debug:
            print(msg)

    def _get_group_name(self, i: int) -> str:
        if self.variant_type == "Animals":
            return self.ANIMALS[i]

        elif self.variant_type == "Letters":
            return self.LETTERS[i]

        elif self.variant_type == "Numbers":
            return str(i + 1)

        else:
            assert False

    def _load_xslx(self):
        return op.load_workbook(str(self.db_file))

    @staticmethod
    def clean_str(dirty_str: str) -> str:
        if dirty_str is not None:
            dirty_str = dirty_str.replace("\u00A0", " ")
            dirty_str = dirty_str.strip()
        return dirty_str

    def get_time_str(self):
        dt = self.time_of_gen.strftime("%d. %m. %Y")
        file_dt = self.time_of_gen.strftime("%Y%m%d%H%M%S")

        return (dt, file_dt)

    def compose_a_doc(self, grp_name: str,
                      is_teacher: bool,
                      rows: list[dict]):
        grp_name = grp_name.capitalize()

        doc = Document()
        category = "Student's Test"
        file_category = "test"
        if is_teacher:
            category = "Teacher's Key"
            file_category = "key"

        dt, file_dt = self.get_time_str()

        doc.add_heading(
            f'{grp_name}: a {category} ___ / {self.db_val_row_cnt*self.headings_cnt-1} pts    Name: ________________', 2)
        if "Instruction" in self.cfg_opts.keys():
            doc.add_paragraph(self.cfg_opts["Instruction"])

        table = doc.add_table(rows=1, cols=self.headings_cnt)
        hdr_cells = table.rows[0].cells
        idx = 0
        for a_heading in self.headings:
            hdr_cells[idx].text = self.headings[idx]
            idx += 1
        for a_row in rows:
            row_cells = table.add_row().cells
            assert row_cells is not None
            idx = 0
            for a_heading in self.headings:
                assert a_row[self.headings[idx]] is not None, f"eq {a_row}"
                row_cells[idx].text = a_row[self.headings[idx]]
                idx += 1

        doc.add_page_break()

        doc.save(str(self.output_dir /
                     f"{grp_name}_{file_category}_{file_dt}.docx"))

    def compose_docx(self, variants: list[dict]):

        for a_group in variants:
            self._l("Composing teacher's key...")
            self.compose_a_doc(a_group, True, variants[a_group]["teacher"])

            self._l("Composing student's test...")
            self.compose_a_doc(a_group, False, variants[a_group]["student"])

    def studentize(self, teachers_rows: list[dict]) -> dict:

        students_rows = deepcopy(teachers_rows)

        idxs = self.headings
        total_cnt = self.headings_cnt
        line_nr = 0
        for a_row in students_rows:
            line_nr += 1
            choice = random.randint(0, total_cnt - 1)
            if a_row[idxs[choice]] is None:
                choice = (choice + 1) % total_cnt
                assert a_row[idxs[choice]
                             ] is not None, f"line: {line_nr}, total {total_cnt} {a_row}"

            for idx in idxs:
                if idx != idxs[choice]:
                    a_row[idx] = ""

        return students_rows

    def _load_cfg_opts(self, cfg_sheet):
        row_cnt = 0
        self.cfg_opts = {}

        for row in cfg_sheet.iter_rows():
            a_row = {}
            a_row_id = 0
            is_empty = True

            if row[0].value is None:
                break
            self.cfg_opts[row[0].value] = row[1].value
            row_cnt += 1
        self._l(f"Loaded {row_cnt} config opts.")

    def _load_db(self, db_sheet):
        row_cnt = 0
        all_rows = []
        self.headings = ()

        for row in db_sheet.iter_rows():
            if row_cnt == 0:
                # Top row.
                for col in row:
                    if col.value is None:
                        break
                    self.headings += (col.value,)
                self.headings_cnt = len(self.headings)
                self._l(f"Headings ({self.headings_cnt}): {self.headings}")
                row_cnt += 1
            else:
                a_row = {}
                a_row_id = 0
                is_empty = True
                for a_col in self.headings:
                    a_row[a_col] = self.clean_str(row[a_row_id].value)
                    a_row_id += 1
                    if a_row[a_col]:
                        is_empty = False
                    else:
                        if not is_empty:
                            a_row[a_col] = ""

                if is_empty:
                    break
                all_rows.append(a_row)
                row_cnt += 1

        self._l(f"Parsed {row_cnt} rows!")
        assert row_cnt >= self.db_val_row_cnt
        return all_rows

    def _generate_variants(self, all_rows):
        variants = {}

        for i in range(self.variant_cnt):
            group_name = self._get_group_name(i)
            self._l(f"Generating variant {group_name}...")

            row_selection = random.sample(all_rows, self.db_val_row_cnt)
            # self._l(json.dumps(row_selection, indent=2))

            random.shuffle(row_selection)
            # self._l(json.dumps(row_selection, indent=2))
            variants[group_name] = {}
            variants[group_name]["teacher"] = row_selection
            variants[group_name]["student"] = self.studentize(row_selection)
        self._l(f"{variants}")
        return variants

    def xform(self):
        wb = self._load_xslx()

        self._load_cfg_opts(wb[self.CFG_SHEET_NAME])
        all_rows = self._load_db(wb[self.DB_SHEET_NAME])
        variants = self._generate_variants(all_rows)
        self.compose_docx(variants)
