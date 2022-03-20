#! /usr/bin/env python3
#! python3
#! /usr/bin/python3
from copy import deepcopy
from datetime import datetime
from os.path import stat
from pathlib import Path
from tkinter import filedialog as tk_fd
from tkinter import ttk
from typing import Tuple
import json
import random

from docx import Document
import openpyxl as op
import tkinter as tk
import tkinter.font as tk_font


class ExcelToTest:

    DB_SHEET_NAME: str = "DB"

    ANIMALS = ["unicorns", "dragons", "ponies",
               "griffins", "hippos", "otters", "bears"]

    LETTERS = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"

    db_file: Path
    variant_type: str
    variant_cnt: int
    verb_cnt: int
    output_dir: Path

    def __init__(self,
                 db_file: Path,
                 variant_type: str,
                 variant_cnt: int,
                 verb_cnt: int,
                 output_dir: Path):

        self.time_of_gen = datetime.now()

        self.db_file = db_file
        assert db_file.is_file()

        self.variant_type = variant_type
        assert variant_type in ["Animals", "Numbers", "Letters"]

        self.variant_cnt = variant_cnt
        if variant_type == "Animals":
            assert variant_cnt <= len(self.ANIMALS)

        if variant_type == "Letters":
            assert variant_cnt <= len(self.LETTERS)

        self.verb_cnt = verb_cnt

        self.output_dir = output_dir
        assert output_dir.is_dir()

    def _get_group_name(self, i: int) -> str:
        if self.variant_type == "Animals":
            return self.ANIMALS[i]

        elif self.variant_type == "Letters":
            return self.LETTERS[i]

        elif self.variant_type == "Numbers":
            return str(i + 1)

        else:
            assert False

    def _load_xslx(self):
        return op.load_workbook(str(self.db_file))

    def _check_format(self, db_sheet):
        assert db_sheet["A1"].value == "Base Form"
        assert db_sheet["B1"].value == "Past Simple"
        assert db_sheet["C1"].value == "Past Participle"
        assert db_sheet["D1"].value == "Czech Translation"

    @staticmethod
    def clean_str(dirty_str: str) -> str:
        if dirty_str is not None:
            dirty_str = dirty_str.replace("\u00A0", " ")
            dirty_str = dirty_str.strip()
        return dirty_str

    def compose_a_doc(self, grp_name: str,
                      is_teacher: bool,
                      verbs: list[dict]):
        grp_name = grp_name.capitalize()

        doc = Document()
        category = "Student's Test"
        file_category = "test"
        if is_teacher:
            category = "Teacher's Key"
            file_category = "key"

        dt = self.time_of_gen.strftime("%-d. %-m. %Y")
        file_dt = self.time_of_gen.strftime("%Y%m%d%H%M%S")

        doc.add_heading(f'{grp_name}: {category} ({dt})', 0)
        doc.add_paragraph(
            'Fill in the empty cells with correct versions and/or translations of the correct irregular verbs.')

        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Basic Form'
        hdr_cells[1].text = 'Past Simple'
        hdr_cells[2].text = 'Past Participle'
        hdr_cells[3].text = 'Czech Translation'
        for a_verb in verbs:
            row_cells = table.add_row().cells
            row_cells[0].text = a_verb["base"]
            row_cells[1].text = a_verb["past_simple"]
            row_cells[2].text = a_verb["past_participle"]
            row_cells[3].text = a_verb["cz"]

        doc.add_page_break()

        doc.save(str(self.output_dir /
                     f"{grp_name}_{file_category}_{file_dt}.docx"))

    def compose_docx(self, variants: list[dict]):

        for a_group in variants:
            print("Composing teacher's key...")
            self.compose_a_doc(a_group, True, variants[a_group]["teacher"])

            print("Composing student's test...")
            self.compose_a_doc(a_group, False, variants[a_group]["student"])

    @staticmethod
    def studentize(teachers_verbs: list[dict]) -> dict:

        students_verbs = deepcopy(teachers_verbs)

        idxs = ["base", "past_simple", "past_participle", "cz"]
        total_cnt = len(idxs)

        for a_verb in students_verbs:
            choice = random.randint(0, total_cnt - 1)
            if a_verb[idxs[choice]] is None:
                choice = (choice + 1) % total_cnt
                assert a_verb[idxs[choice]] is not None

            for idx in idxs:
                if idx != idxs[choice]:
                    a_verb[idx] = ""

        return students_verbs

    def xform(self):
        wb = self._load_xslx()

        db_sheet = wb[self.DB_SHEET_NAME]
        self._check_format(db_sheet)

        row_cnt = 0

        all_verbs = []

        for row in db_sheet.iter_rows():
            if row_cnt != 0:
                a_verb = {
                    "base": self.clean_str(row[0].value),
                    "past_simple": self.clean_str(row[1].value),
                    "past_participle": self.clean_str(row[2].value),
                    "cz": self.clean_str(row[3].value)
                }
                all_verbs.append(a_verb)
            row_cnt += 1

        print(f"Parsed {row_cnt} rows!")
        assert row_cnt >= self.verb_cnt

        # print(json.dumps(all_verbs, indent=2))

        variants = {}

        for i in range(self.variant_cnt):
            group_name = self._get_group_name(i)
            print(f"Generating variant {group_name}...")

            verb_selection = random.sample(all_verbs, self.verb_cnt)
            # print(json.dumps(verb_selection, indent=2))

            random.shuffle(verb_selection)
            # print(json.dumps(verb_selection, indent=2))
            variants[group_name] = {}
            variants[group_name]["teacher"] = verb_selection
            variants[group_name]["student"] = self.studentize(verb_selection)

        print(f"{variants}")
        self.compose_docx(variants)


class Style:
    pad_x: int
    pad_y: int
    mar_x: int
    mar_y: int
    fonts: dict

    def __init__(self,
                 pad_x: int = 2,
                 pad_y: int = 2,
                 mar_x: int = 1,
                 mar_y: int = 1):
        self.pad_x = pad_x
        self.pad_y = pad_y
        self.mar_x = mar_x
        self.mar_y = mar_y
        self.fonts = {}

    def set_font(self,
                 cls: str="",
                 family: str="Arial",
                 size: int=10,
                 weight: str="normal",
                 slant: str="roman"):
        self.fonts[cls] = tk_font.Font(family=family,
                                       size=size,
                                       weight=weight,
                                       slant=slant)

    def get_font(self, cls=""):
        return self.fonts[cls]


class IntField(tk.Entry):

    var: tk.IntVar
    master: tk.Widget
    root: tk.Widget
    label: str

    def __init__(self,
                 master: tk.Widget,
                 label: str,
                 value: int,
                 **kwargs) -> None:
        self.master = master
        pi = master.winfo_parent()
        self.root = master._nametowidget(pi)
        self.label = label
        self.var = tk.IntVar(value=value)
        super().__init__(master, kwargs)
        self.config(textvariable=self.var)

    def draw(self, r: int=0, c: int=0) -> (int, int):

        label = tk.Label(self.master,
                         text=self.label,
                         padx=self.root.s.pad_x,
                         pady=self.root.s.pad_y,
                         font=self.root.s.get_font("regular"))
        label.grid(row=r, column=c, sticky="E")

        self.grid(row=r, column=c + 1, sticky="W")

        return (r + 1, c + 2)

    def get_value(self) -> int:
        return self.var.get()


class PickField(tk.OptionMenu):
    var: tk.StringVar
    master: tk.Widget
    root: tk.Widget
    label: str

    def __init__(self,
                 master: tk.Widget,
                 label: str,
                 picks: Tuple[str]) -> None:
        self.master = master
        pi = master.winfo_parent()
        self.root = master._nametowidget(pi)
        self.var = tk.StringVar(value=picks[0])
        self.picks = picks
        self.label = label
        super().__init__(master,
                         self.var,
                         *picks)

    def draw(self, r: int=0, c: int=0) -> (int, int):

        label = tk.Label(self.master,
                         text=self.label,
                         padx=self.root.s.pad_x,
                         pady=self.root.s.pad_y,
                         font=self.root.s.get_font("regular"))
        label.grid(row=r, column=c, sticky="E")
        self.grid(row=r, column=c + 1, sticky="W")

        return (r + 1, c + 2)

    def get_value(self) -> int:
        return self.var.get()


class OpenDialogField(tk.Entry):

    var: tk.StringVar
    master: tk.Widget
    root: tk.Widget
    label: str
    last_dir: Path

    def __init__(self,
                 master: tk.Widget,
                 label: str,
                 value: str,
                 is_file: bool=True,
                 **kwargs) -> None:
        self.master = master
        pi = master.winfo_parent()
        self.root = master._nametowidget(pi)
        self.label = label
        self.var = tk.StringVar(value=value)
        self.last_dir = Path.cwd()
        self.last_filename = None
        if not is_file:
            self.last_filename = self.last_dir
        self.is_file = is_file
        super().__init__(master, kwargs)
        self.config(textvariable=self.var)

    def open_file(self):
        filetypes = [('Database Files', '*.xlsx')]

        filename = \
            tk_fd.askopenfilename(title='Open a database .xlsx file',
                                  initialdir=self.last_dir,
                                  filetypes=filetypes)
        filename_path = Path(filename)

        if filename_path.is_file():
            self.last_filename = filename_path
            self.last_dir = self.last_filename.parent
            self.var.set(str(self.last_filename))
        else:
            tk.messagebox.showerror(title="Error",
                                    message="Database file invalid")

    def open_dir(self):
        dirname = \
            tk_fd.askdirectory(title='Select output directory',
                               initialdir=self.last_dir)
        dirname_path = Path(dirname)

        if dirname_path.is_dir():
            self.last_filename = dirname_path
            self.last_dir = self.last_filename.parent
            self.var.set(str(self.last_filename))
        else:
            tk.messagebox.showerror(title="Error",
                                    message="Output folder invalid")

    def btn_click(self):

        if self.is_file:
            self.open_file()
        else:
            self.open_dir()

    def draw(self, r: int=0, c: int=0) -> (int, int):

        label = tk.Label(self.master,
                         text=self.label,
                         padx=self.root.s.pad_x,
                         pady=self.root.s.pad_y,
                         font=self.root.s.get_font("regular"))

        btn = tk.Button(self.master, text="Open", command=self.btn_click)

        label.grid(row=r,
                   column=c,
                   columnspan=2,
                   sticky="W")
        self.grid(row=r + 1, column=c,
                  sticky="E")
        btn.grid(row=r + 1, column=c + 1, sticky="W")

        return (r + 2, c + 2)

    def get_value(self):
        return self.last_filename


class IrregularVerbs(tk.Tk):

    xlsx_db_file_fld: OpenDialogField

    output_dir_fld: OpenDialogField

    variant_cnt_fld: IntField

    verb_cnt_fld: IntField

    variant_type_fld: PickField

    s: Style

    def __init__(self):
        super().__init__()

        self.s = Style()
        self.s.set_font("regular")

        self.main_frame = tk.Frame(self)

        self.xlsx_db_file_fld = OpenDialogField(self.main_frame,
                                                label="Database (.xlsx)",
                                                value="",
                                                width=20)

        self.output_dir_fld = OpenDialogField(self.main_frame,
                                              label="Output folder",
                                              value=str(Path.cwd()),
                                              is_file=False,
                                              width=20)

        self.variant_type_fld = PickField(self.main_frame,
                                          label="Variant Type",
                                          picks=(
                                              "Animals", "Numbers", "Letters"))

        self.variant_cnt_fld = IntField(self.main_frame,
                                        label="Variant Count",
                                        value=2,
                                        width=20)

        self.verb_cnt_fld = IntField(self.main_frame,
                                     label="Verb Count",
                                     value=5,
                                     width=20)

        self.generate_btn = tk.Button(self.main_frame,
                                      text="Generate",
                                      command=self.generate_variants)

    def generate_variants(self):
        ett = ExcelToTest(db_file=self.xlsx_db_file_fld.get_value(),
                          variant_type=self.variant_type_fld.get_value(),
                          variant_cnt=self.variant_cnt_fld.get_value(),
                          verb_cnt=self.verb_cnt_fld.get_value(),
                          output_dir=self.output_dir_fld.get_value())

        ett.xform()

    def draw(self):
        r = 0
        r, _ = self.xlsx_db_file_fld.draw(r, 0)
        r, _ = self.output_dir_fld.draw(r, 0)
        r, _ = self.variant_type_fld.draw(r, 0)
        r, _ = self.variant_cnt_fld.draw(r, 0)
        r, _ = self.verb_cnt_fld.draw(r, 0)
        self.generate_btn.grid(row=r,
                               column=0,
                               columnspan=2,
                               sticky="EW")

        self.main_frame.grid(sticky="NW")

        self.geometry("400x240")

        return (r, 2)


if __name__ == "__main__":
    # Main

    window = IrregularVerbs()
    window.draw()
    window.mainloop()
